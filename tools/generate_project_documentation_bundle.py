from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"

SOURCE_FILES = [
    DOCS_DIR / "00_EXECUTIVE_SUMMARY.md",
    DOCS_DIR / "01_TOOLS_TECH_STACK.md",
    DOCS_DIR / "02_WORKFLOWS.md",
    DOCS_DIR / "03_DATA_MODELS.md",
    DOCS_DIR / "05_CODE_WALKTHROUGH.md",
    DOCS_DIR / "06_DEPLOYMENT.md",
    DOCS_DIR / "07_DETAILED_CODE_GUIDE.md",
    DOCS_DIR / "FULL_DOCUMENTATION.md",
]

OUTPUT_DOCX = DOCS_DIR / "Cartiqo_Project_Documentation_Updated.docx"
OUTPUT_PDF = DOCS_DIR / "Cartiqo_Project_Documentation_Updated.pdf"


def iter_markdown_lines(paths: Iterable[Path]) -> list[str]:
    lines: list[str] = []
    for path in paths:
        lines.append(f"# {path.stem.replace('_', ' ')}")
        lines.extend(path.read_text(encoding="utf-8").splitlines())
        lines.append("")
    return lines


def build_docx(lines: list[str]) -> None:
    document = Document()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cartiqo / SwiftCart Project Documentation")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Updated Azure-based system documentation").italic = True

    document.add_paragraph("")
    in_code_block = False
    code_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code_block:
                paragraph = document.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines.clear()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line:
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:3].isdigit() and line[1:3] == ". ":
            document.add_paragraph(line[3:].strip(), style="List Number")
        else:
            document.add_paragraph(line)

    document.save(OUTPUT_DOCX)


def build_pdf(lines: list[str]) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCenter",
        parent=styles["Title"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0F172A"),
        fontSize=20,
        leading=24,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=10,
        spaceAfter=6,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=8,
        spaceAfter=4,
    )
    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=6,
        spaceAfter=3,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#111827"),
    )
    code = ParagraphStyle(
        "Code",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8.5,
        leading=10.5,
        backColor=colors.HexColor("#F3F4F6"),
    )

    story = [
        Paragraph("Cartiqo / SwiftCart Project Documentation", title_style),
        Paragraph("Updated Azure-based system documentation", body),
        Spacer(1, 12),
    ]

    in_code_block = False
    code_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code_block:
                story.append(Preformatted("\n".join(code_lines), code))
                story.append(Spacer(1, 6))
                code_lines.clear()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line:
            story.append(Spacer(1, 4))
            continue

        if line.startswith("# "):
            story.append(Paragraph(line[2:].strip(), h1))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:].strip(), h2))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), h3))
        elif line.startswith("- "):
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(line[2:].strip(), body))],
                    bulletType="bullet",
                    leftIndent=16,
                )
            )
        elif line[:3].isdigit() and line[1:3] == ". ":
            story.append(Paragraph(line, body))
        else:
            story.append(Paragraph(line, body))

    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=30,
        bottomMargin=30,
    )
    doc.build(story)


def main() -> None:
    lines = iter_markdown_lines(SOURCE_FILES)
    build_docx(lines)
    build_pdf(lines)
    print(OUTPUT_DOCX)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    main()

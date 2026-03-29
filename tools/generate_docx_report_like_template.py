from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\GS\OneDrive\Documents\SUNNY  DOC2.docx")
OUT_DOCX = REPO_ROOT / "docs" / "Cartiqo_Detailed_Report.docx"

LIB_DIR = REPO_ROOT / "lib"
FUNCTIONS_INDEX = REPO_ROOT / "functions" / "src" / "index.ts"
PUBSPEC = REPO_ROOT / "pubspec.yaml"


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _clear_doc_body(doc: Document) -> None:
    body = doc._element.body  # noqa: SLF001 (python-docx internal)
    # Preserve the section properties element if present.
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
            break

    for child in list(body):
        if sect_pr is not None and child is sect_pr:
            continue
        body.remove(child)

    if sect_pr is not None and sect_pr.getparent() is None:
        body.append(sect_pr)


def _h(doc: Document, text: str, level: int) -> None:
    style = {1: "Heading2", 2: "Heading4", 3: "Heading6"}.get(level, "Heading6")
    p = doc.add_paragraph(text, style=style)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _p(doc: Document, text: str, style: str = "BodyText") -> None:
    doc.add_paragraph(text, style=style)


def _bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="ListParagraph")


def _extract_dependencies() -> list[str]:
    txt = _read_text(PUBSPEC)
    deps = []
    in_deps = False
    for line in txt.splitlines():
        if line.strip() == "dependencies:":
            in_deps = True
            continue
        if in_deps and re.match(r"^[a-zA-Z_]", line):
            # next top-level section
            break
        if in_deps:
            m = re.match(r"^\s{2}([a-zA-Z0-9_]+):\s*(.*)$", line)
            if m:
                name, ver = m.group(1), m.group(2).strip()
                if name == "flutter":
                    continue
                deps.append(f"{name} {ver}".rstrip())
    return deps


def _extract_firestore_collections() -> list[str]:
    pattern = re.compile(r"\.collection(?:Group)?\('([^']+)'\)")
    collections = set()
    for path in list(LIB_DIR.rglob("*.dart")) + ([FUNCTIONS_INDEX] if FUNCTIONS_INDEX.exists() else []):
        txt = _read_text(path)
        for m in pattern.finditer(txt):
            collections.add(m.group(1))
    return sorted(collections)


def _extract_screens() -> list[str]:
    out = []
    for p in (LIB_DIR / "screens").glob("*.dart"):
        src = _read_text(p)
        classes = re.findall(r"^\s*class\s+([A-Za-z0-9_]+)\b", src, flags=re.M)
        if classes:
            out.append(f"{p.name}: {', '.join(classes)}")
    return sorted(out)


def _extract_providers() -> list[str]:
    out = []
    for p in (LIB_DIR / "providers").glob("*.dart"):
        src = _read_text(p)
        classes = re.findall(r"^\s*class\s+([A-Za-z0-9_]+)\b", src, flags=re.M)
        if classes:
            out.append(f"{p.name}: {', '.join(classes)}")
    return sorted(out)


def _extract_functions_exports() -> list[str]:
    if not FUNCTIONS_INDEX.exists():
        return []
    src = _read_text(FUNCTIONS_INDEX)
    return re.findall(r"export\s+const\s+([A-Za-z0-9_]+)\s*=", src)


def generate() -> Path:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")

    doc = Document(str(TEMPLATE))
    _clear_doc_body(doc)

    stamp = datetime.now().strftime("%d-%m-%Y")
    title = doc.add_paragraph("CARTIQO (SWIFTCART) MINI PROJECT REPORT", style="Heading2")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p(doc, f"Generated on: {stamp}", style="BodyText")
    doc.add_paragraph("")

    # INDEX (matches the template’s style naming)
    _h(doc, "INDEX", level=1)
    _bullets(
        doc,
        [
            "CHAPTER-1 INTRODUCTION",
            "CHAPTER-2 BACKGROUND AND LITERATURE",
            "CHAPTER-3 REQUIREMENTS SPECIFICATION",
            "CHAPTER-4 SYSTEM STUDY AND DESIGN",
            "CHAPTER-5 DEVELOPMENT TOOLS",
            "CHAPTER-6 IMPLEMENTATION (CODE WALKTHROUGH)",
            "CHAPTER-7 DATABASE DESIGN (FIRESTORE)",
            "CHAPTER-8 TESTING AND DEPLOYMENT",
            "CHAPTER-9 CONCLUSION AND FUTURE ENHANCEMENTS",
        ],
    )

    doc.add_page_break()

    # CHAPTER-1
    _h(doc, "CHAPTER-1 INTRODUCTION", level=1)
    _h(doc, "PROJECT OVERVIEW", level=1)
    _p(
        doc,
        "Cartiqo is a Flutter + Firebase retail checkout and mall-operations platform. "
        "It provides a customer-facing scan-to-bill workflow and a web portal for admins and mall managers to manage inventory, barcodes, billing settings, subscriptions, and sales analytics.",
    )
    _h(doc, "OBJECTIVE", level=1)
    _bullets(
        doc,
        [
            "Reduce checkout time using barcode-first scanning and consistent billing rules.",
            "Centralize mall operations: subscription management, manager accounts, inventory, labels, and reporting.",
            "Enable export-ready reporting (CSV) and print-ready artifacts (PDF) from the portal.",
        ],
    )

    doc.add_page_break()

    # CHAPTER-2
    _h(doc, "CHAPTER-2 BACKGROUND AND LITERATURE", level=1)
    _p(
        doc,
        "This chapter summarizes the technologies used and why they are suitable for building a cross-platform retail operations system.",
    )
    _h(doc, "Flutter and Cross-Platform Development", level=2)
    _p(
        doc,
        "Flutter provides a single codebase for Android, iOS, and Web, with a reactive UI framework and strong performance characteristics.",
    )
    _h(doc, "Firebase (Auth, Firestore, Hosting)", level=2)
    _p(
        doc,
        "Firebase offers managed authentication, a scalable NoSQL database (Firestore), and reliable hosting for web deployments.",
    )
    _h(doc, "Barcode/QR Operations", level=2)
    _p(
        doc,
        "Barcode scanning and QR generation are used for fast product identification and label workflows in retail environments.",
    )

    doc.add_page_break()

    # CHAPTER-3
    _h(doc, "CHAPTER-3 REQUIREMENTS SPECIFICATION", level=1)
    _h(doc, "FUNCTIONAL REQUIREMENTS", level=2)
    _bullets(
        doc,
        [
            "User authentication (email/password, phone OTP).",
            "Mall selection and mall-scoped product catalog browsing/search.",
            "Barcode scanning to add items to cart and compute billing totals.",
            "Admin management of malls, subscriptions, and manager accounts.",
            "Mall manager inventory management, barcode mapping, and billing settings.",
            "Sales history dashboard with filtering, grouping, refresh, and CSV export.",
        ],
    )
    _h(doc, "NON-FUNCTIONAL REQUIREMENTS", level=2)
    _bullets(
        doc,
        [
            "Fast UI responsiveness during scanning and checkout.",
            "Data consistency and auditability for operational actions.",
            "Secure handling of credentials and secrets (no keys committed to repo).",
            "Scalable database access patterns (indexed queries where required).",
        ],
    )

    doc.add_page_break()

    # CHAPTER-4
    _h(doc, "CHAPTER-4 SYSTEM STUDY AND DESIGN", level=1)
    _h(doc, "SYSTEM ARCHITECTURE", level=2)
    _bullets(
        doc,
        [
            "Client: Flutter app (mobile + web).",
            "Backend: Firebase Auth + Cloud Firestore.",
            "Web deployment: Firebase Hosting.",
            "Optional backend automation: Firebase Cloud Functions (OTP + cleanup).",
        ],
    )
    _h(doc, "USER WORKFLOWS (HIGH LEVEL)", level=2)
    _bullets(
        doc,
        [
            "Customer: login → select mall → scan items → cart → checkout → history.",
            "Admin (web): login → manage malls/subscriptions → manage managers → reports/exports.",
            "Manager (web): login → inventory/barcodes → billing settings → sales dashboard → export CSV.",
        ],
    )
    _h(doc, "UI MODULES (SCREENS)", level=2)
    _bullets(doc, _extract_screens())

    doc.add_page_break()

    # CHAPTER-5
    _h(doc, "CHAPTER-5 DEVELOPMENT TOOLS", level=1)
    _h(doc, "SYSTEM ENVIRONMENT", level=2)
    _bullets(
        doc,
        [
            "Flutter SDK + Dart",
            "Firebase CLI",
            "Node.js (for Cloud Functions)",
            "Chrome (web testing) and Android device/emulator (mobile testing)",
        ],
    )
    _h(doc, "KEY PACKAGES (pubspec.yaml)", level=2)
    _bullets(doc, _extract_dependencies())

    doc.add_page_break()

    # CHAPTER-6
    _h(doc, "CHAPTER-6 IMPLEMENTATION", level=1)
    _h(doc, "STATE MANAGEMENT (PROVIDERS)", level=2)
    _bullets(doc, _extract_providers())
    _h(doc, "SERVICES", level=2)
    _bullets(
        doc,
        [
            "CSV export service (subscription and sales dashboard export).",
            "PDF/printing services (labels/receipts).",
            "File download helper for web exports.",
        ],
    )

    doc.add_page_break()

    # CHAPTER-7
    _h(doc, "CHAPTER-7 DATABASE DESIGN (FIRESTORE)", level=1)
    _h(doc, "COLLECTIONS USED", level=2)
    _bullets(doc, _extract_firestore_collections())
    _h(doc, "KEY SUBCOLLECTION PATTERNS", level=2)
    _bullets(
        doc,
        [
            "malls/{mallId}/products",
            "malls/{mallId}/barcodes",
            "malls/{mallId}/payments",
            "malls/{mallId}/managers",
            "malls/{mallId}/promotions",
            "malls/{mallId}/staff_activity",
            "users/{uid}/bills",
            "users/{uid}/payments",
        ],
    )

    doc.add_page_break()

    # CHAPTER-8
    _h(doc, "CHAPTER-8 TESTING AND DEPLOYMENT", level=1)
    _h(doc, "WEB DEPLOYMENT (FIREBASE HOSTING)", level=2)
    _p(doc, "Build and deploy:", style="BodyText")
    _bullets(
        doc,
        [
            "flutter build web --release",
            "firebase deploy --only hosting",
        ],
    )
    _h(doc, "CLOUD FUNCTIONS (OPTIONAL)", level=2)
    exports = _extract_functions_exports()
    if exports:
        _p(doc, "Functions implemented in functions/src/index.ts:", style="BodyText")
        _bullets(doc, exports)
    _p(
        doc,
        "Note: Deploying Cloud Functions may require enabling Cloud Build / Artifact Registry APIs, which can require a Blaze plan on Firebase.",
        style="BodyText",
    )

    doc.add_page_break()

    # CHAPTER-9
    _h(doc, "CHAPTER-9 CONCLUSION", level=1)
    _p(
        doc,
        "Cartiqo provides a unified scan-to-bill customer experience and an operational portal for mall teams. "
        "The architecture is designed around Firebase-managed services and a mall-scoped Firestore schema, enabling rapid iteration and scalable operations.",
        style="BodyText",
    )
    _h(doc, "FUTURE ENHANCEMENTS", level=2)
    _bullets(
        doc,
        [
            "Harden Firestore security rules with role-based access control (RBAC).",
            "Automated settlement/reconciliation reporting and alerts.",
            "More advanced analytics (cohorts, category trends, forecasting).",
            "Formal audit logs and admin action trails for compliance-heavy deployments.",
        ],
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


if __name__ == "__main__":
    out = generate()
    print(out)
